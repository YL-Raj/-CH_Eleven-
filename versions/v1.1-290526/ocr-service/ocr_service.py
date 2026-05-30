# Copyright (c) 2026 RAJ.Y — All rights reserved.
# CH_Eleven Fantasy Cricket Platform
# https://github.com/RAJ-Y/ch-eleven
"""
OCR Service - Flask microservice for PDF/Image scorecard extraction
Handles PDF uploads, OCR, text parsing, and JSON generation
Production-ready with proper error handling, logging, and security
"""

import os
import json
import logging
import tempfile
import shutil
from datetime import datetime
from functools import wraps
from pathlib import Path

from flask import Flask, request, jsonify
from flask_cors import CORS
from werkzeug.utils import secure_filename
from pdf2image import convert_from_path
from docx import Document
import pytesseract
from PIL import Image

from scorecard_parser import ScorecardParser
from match_data_parser import MatchDataParser

# ─── Configuration ──────────────────────────────────────────────────────────
app = Flask(__name__)
CORS(app)

# Logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Config
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
ALLOWED_EXTENSIONS = {'pdf', 'png', 'jpg', 'jpeg', 'docx', 'csv'}
TEMP_DIR = Path('/tmp/uploads')
TEMP_DIR.mkdir(exist_ok=True)

# Initialize parsers
parser = ScorecardParser()
match_parser = MatchDataParser()

# ─── Utilities ──────────────────────────────────────────────────────────────
def require_auth(f):
    """Decorator to check admin key"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        admin_key = request.headers.get('x-admin-key')
        expected_key = os.getenv('ADMIN_SECRET', 'ch11-admin-2026')
        
        if not admin_key or admin_key != expected_key:
            logger.warning(f"Unauthorized OCR upload attempt from {request.remote_addr}")
            return jsonify({'error': 'Unauthorized'}), 401
        
        return f(*args, **kwargs)
    return decorated_function

def allowed_file(filename: str) -> bool:
    """Check if file extension is allowed"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def cleanup_file(filepath: str):
    """Securely delete temporary file"""
    try:
        if os.path.exists(filepath):
            os.remove(filepath)
            logger.info(f"Cleaned up: {filepath}")
    except Exception as e:
        logger.error(f"Cleanup error: {e}")

# ─── File Extraction ────────────────────────────────────────────────────────
def extract_pdf(filepath: str) -> str:
    """Extract text from PDF using OCR"""
    logger.info(f"Extracting PDF: {filepath}")
    text = ""
    try:
        # Convert PDF to images
        images = convert_from_path(filepath, dpi=150)
        
        for i, image in enumerate(images):
            logger.info(f"OCR processing page {i+1}/{len(images)}")
            # Run Tesseract OCR
            page_text = pytesseract.image_to_string(image)
            text += f"\n--- PAGE {i+1} ---\n" + page_text
        
        return text
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        raise

def extract_image(filepath: str) -> str:
    """Extract text from image using OCR"""
    logger.info(f"Extracting image: {filepath}")
    try:
        image = Image.open(filepath)
        text = pytesseract.image_to_string(image)
        return text
    except Exception as e:
        logger.error(f"Image extraction error: {e}")
        raise

def extract_docx(filepath: str) -> str:
    """Extract text from DOCX"""
    logger.info(f"Extracting DOCX: {filepath}")
    text = ""
    try:
        doc = Document(filepath)
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
        return text
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        raise

def extract_csv(filepath: str) -> str:
    """Extract text from CSV"""
    logger.info(f"Extracting CSV: {filepath}")
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        logger.error(f"CSV extraction error: {e}")
        raise

# ─── Routes ──────────────────────────────────────────────────────────────────
@app.route('/health', methods=['GET'])
def health():
    """Health check endpoint"""
    return jsonify({'status': 'healthy', 'service': 'ocr-service'}), 200

@app.route('/api/ocr/extract', methods=['POST'])
@require_auth
def ocr_extract():
    """
    Upload scorecard file (PDF/Image/DOCX/CSV) and extract text
    
    Returns:
        JSON with extracted text and parsing results
    """
    try:
        # Validate request
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400

        file = request.files['file']
        inning = request.form.get('inning', '1', type=int)

        if not file or file.filename == '':
            return jsonify({'error': 'Invalid file'}), 400

        if not allowed_file(file.filename):
            return jsonify({'error': f'File type not allowed. Allowed: {", ".join(ALLOWED_EXTENSIONS)}'}), 400

        if file.content_length and file.content_length > MAX_FILE_SIZE:
            return jsonify({'error': f'File too large. Max: {MAX_FILE_SIZE / 1024 / 1024}MB'}), 413

        # Save temporarily
        filename = secure_filename(file.filename)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S_')
        temp_filename = f"{timestamp}{filename}"
        temp_filepath = TEMP_DIR / temp_filename

        file.save(str(temp_filepath))
        logger.info(f"File saved: {temp_filename}")

        try:
            # Extract text based on file type
            ext = filename.rsplit('.', 1)[1].lower()
            
            if ext == 'pdf':
                raw_text = extract_pdf(str(temp_filepath))
            elif ext in ['png', 'jpg', 'jpeg']:
                raw_text = extract_image(str(temp_filepath))
            elif ext == 'docx':
                raw_text = extract_docx(str(temp_filepath))
            elif ext == 'csv':
                raw_text = extract_csv(str(temp_filepath))
            else:
                return jsonify({'error': 'Unsupported file type'}), 400

            # Parse scorecard
            scorecard = parser.parse_text(raw_text, inning=inning)

            logger.info(f"Extraction complete. Batting: {len(scorecard['batting'])}, Bowling: {len(scorecard['bowling'])}")

            return jsonify({
                'success': True,
                'filename': filename,
                'inning': inning,
                'extracted_text_preview': raw_text[:500],  # First 500 chars
                'scorecard': scorecard,
                'timestamp': datetime.now().isoformat(),
            }), 200

        finally:
            # Cleanup temp file
            cleanup_file(str(temp_filepath))

    except Exception as e:
        logger.error(f"OCR extraction error: {e}", exc_info=True)
        return jsonify({'error': f'Processing error: {str(e)}'}), 500

@app.route('/api/ocr/batch', methods=['POST'])
@require_auth
def ocr_batch():
    """
    Process multiple files (1st and 2nd inning)
    
    Expected: files with inning numbers
    """
    try:
        results = []
        
        if 'files' not in request.files:
            return jsonify({'error': 'No files provided'}), 400

        files = request.files.getlist('files')
        innings = request.form.getlist('innings', type=int)

        if len(files) > 2:
            return jsonify({'error': 'Max 2 files allowed (1st and 2nd inning)'}), 400

        for i, file in enumerate(files):
            if not file or file.filename == '':
                continue

            inning = innings[i] if i < len(innings) else (i + 1)

            # Save temporarily
            filename = secure_filename(file.filename)
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S_')
            temp_filename = f"{timestamp}{filename}"
            temp_filepath = TEMP_DIR / temp_filename

            file.save(str(temp_filepath))

            try:
                # Extract
                ext = filename.rsplit('.', 1)[1].lower()
                
                if ext == 'pdf':
                    raw_text = extract_pdf(str(temp_filepath))
                elif ext in ['png', 'jpg', 'jpeg']:
                    raw_text = extract_image(str(temp_filepath))
                elif ext == 'docx':
                    raw_text = extract_docx(str(temp_filepath))
                elif ext == 'csv':
                    raw_text = extract_csv(str(temp_filepath))
                else:
                    continue

                # Parse
                scorecard = parser.parse_text(raw_text, inning=inning)
                results.append({
                    'inning': inning,
                    'filename': filename,
                    'scorecard': scorecard,
                })

            finally:
                cleanup_file(str(temp_filepath))

        return jsonify({
            'success': True,
            'processed': len(results),
            'scorecards': results,
            'timestamp': datetime.now().isoformat(),
        }), 200

    except Exception as e:
        logger.error(f"Batch processing error: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500

@app.route('/api/ocr/status', methods=['GET'])
def status():
    """Service status endpoint"""
    try:
        # Check temp directory
        temp_files = list(TEMP_DIR.glob('*'))
        
        return jsonify({
            'status': 'operational',
            'temp_files': len(temp_files),
            'tesseract_available': True,
            'max_file_size_mb': MAX_FILE_SIZE / 1024 / 1024,
            'supported_formats': list(ALLOWED_EXTENSIONS),
        }), 200
    except Exception as e:
        logger.error(f"Status check error: {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/api/csv/parse', methods=['POST'])
@require_auth
def csv_parse():
    """
    Parse a CSV (or JSON) match data file and return structured match data.
    Accepts form field 'csv_data' (raw text) or a 'file' upload.
    """
    try:
        # Try form field first (sent by the Node backend as form-data)
        csv_data = request.form.get('csv_data')
        file_type = request.form.get('format', 'csv')

        if not csv_data and 'file' in request.files:
            f = request.files['file']
            csv_data = f.read().decode('utf-8')
            ext = f.filename.rsplit('.', 1)[-1].lower() if f.filename else 'csv'
            file_type = ext if ext in ('csv', 'json') else 'csv'

        if not csv_data:
            return jsonify({'error': 'No data provided'}), 400

        result = match_parser.parse_file(csv_data, file_type)

        if 'error' in result:
            return jsonify({'error': result['error']}), 400

        return jsonify(result), 200

    except Exception as e:
        logger.error(f"CSV parse error: {e}", exc_info=True)
        return jsonify({'error': f'Parse error: {str(e)}'}), 500

# ─── Error Handlers ─────────────────────────────────────────────────────────
@app.errorhandler(413)
def request_entity_too_large(error):
    return jsonify({'error': 'File too large'}), 413

@app.errorhandler(404)
def not_found(error):
    return jsonify({'error': 'Not found'}), 404

@app.errorhandler(500)
def internal_error(error):
    logger.error(f"Internal error: {error}")
    return jsonify({'error': 'Internal server error'}), 500

# ─── Cleanup on Startup ─────────────────────────────────────────────────────
@app.before_request
def cleanup_old_files():
    """Remove files older than 1 hour"""
    try:
        now = datetime.now().timestamp()
        for file in TEMP_DIR.glob('*'):
            if os.path.isfile(file):
                if (now - os.path.getmtime(file)) > 3600:  # 1 hour
                    cleanup_file(str(file))
    except Exception as e:
        logger.warning(f"Cleanup error: {e}")

if __name__ == '__main__':
    # Use gunicorn in production
    port = int(os.getenv('OCR_PORT', 5000))
    debug = os.getenv('DEBUG', 'False').lower() == 'true'
    
    logger.info(f"Starting OCR service on port {port}")
    app.run(host='0.0.0.0', port=port, debug=debug, threaded=True)
